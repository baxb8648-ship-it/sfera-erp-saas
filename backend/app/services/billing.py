import os
import uuid
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from sqlalchemy.orm import Session

from ..models import Tenant, Invoice, Organization

# Реквизиты по умолчанию для владельца платформы СФЕРУМ (Поставщик)
PROVIDER_REQUISITES = {
    "name": "ООО 'СФЕРУМ ТЕХНОЛОГИИ'",
    "inn": "5610248560",
    "kpp": "561001001",
    "address": "Россия, Оренбургская обл., г. Оренбург, ул. Мира, д. 12",
    "bank_name": "АКБ 'АВАНГАРД' - ПАО, Г. МОСКВА",
    "bik": "044525201",
    "rs": "40702810404500001234",
    "ks": "30101810000000000201"
}


def num2str_ru(val: float) -> str:
    """Простая функция конвертации суммы в строку прописью (для B2B счетов)"""
    # Для простоты и исключения внешних зависимостей сделаем лаконичную B2B версию
    rub = int(val)
    kop = int(round((val - rub) * 100))
    
    # Минимальный маппинг числительных для счетов до 100 тыс. руб.
    units = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    
    if rub == 0:
        return f"Ноль рублей {kop:02d} копеек"
        
    parts = []
    
    # Разбираем тысячи (до 99 тысяч)
    thousands = rub // 1000
    rem = rub % 1000
    
    if thousands > 0:
        th_tens = thousands // 10
        th_units = thousands % 10
        if th_tens == 1:
            parts.append(teens[th_units])
            parts.append("тысяч")
        else:
            if th_tens > 1:
                parts.append(tens[th_tens])
            if th_units == 1:
                parts.append("одна")
                parts.append("тысяча")
            elif th_units in [2, 3, 4]:
                parts.append(units[th_units])
                parts.append("тысячи")
            else:
                if th_units > 0:
                    parts.append(units[th_units])
                parts.append("тысяч")
                
    # Разбираем сотни, десятки, единицы
    h = rem // 100
    t = (rem % 100) // 10
    u = rem % 10
    
    if h > 0:
        parts.append(hundreds[h])
    if t == 1:
        parts.append(teens[u])
    else:
        if t > 1:
            parts.append(tens[t])
        if u > 0:
            if u == 1:
                parts.append("один")
            elif u == 2:
                parts.append("два")
            else:
                parts.append(units[u])
                
    # Склонение слова "рубль"
    last_digit = rub % 10
    last_two = rub % 100
    if last_two in [11, 12, 13, 14]:
        parts.append("рублей")
    elif last_digit == 1:
        parts.append("рубль")
    elif last_digit in [2, 3, 4]:
        parts.append("рубля")
    else:
        parts.append("рублей")
        
    parts_str = " ".join([p for p in parts if p])
    # Делаем первую букву заглавной
    parts_str = parts_str[0].upper() + parts_str[1:]
    
    return f"{parts_str} {kop:02d} копеек"


def generate_invoice_docx(invoice_id: int, tenant: Tenant, amount: float, output_dir: str) -> str:
    """Генерирует профессиональный B2B счет на оплату в формате Docx"""
    doc = Document()
    
    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        
    # Стили шрифтов
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # 1. Шапка со сведениями о банке получателя
    table_bank = doc.add_table(rows=4, cols=4)
    table_bank.style = 'Table Grid'
    
    # Заполняем реквизиты банка поставщика (СФЕРУМ)
    hdr_cells = table_bank.rows[0].cells
    hdr_cells[0].text = PROVIDER_REQUISITES["bank_name"]
    hdr_cells[2].text = "БИК"
    hdr_cells[3].text = PROVIDER_REQUISITES["bik"]
    
    row1 = table_bank.rows[1].cells
    row1[0].text = "Банк получателя"
    row1[2].text = "Сч. №"
    row1[3].text = PROVIDER_REQUISITES["ks"]
    
    row2 = table_bank.rows[2].cells
    row2[0].text = f"ИНН {PROVIDER_REQUISITES['inn']}"
    row2[1].text = f"КПП {PROVIDER_REQUISITES['kpp']}"
    row2[2].text = "Сч. №"
    row2[3].text = PROVIDER_REQUISITES["rs"]
    
    row3 = table_bank.rows[3].cells
    row3[0].text = PROVIDER_REQUISITES["name"]
    row3[0].paragraphs[0].runs[0].font.bold = True
    row3[2].text = ""
    row3[3].text = ""
    
    # Объединение ячеек в шапке банка для красоты
    table_bank.rows[0].cells[0].merge(table_bank.rows[0].cells[1])
    table_bank.rows[1].cells[0].merge(table_bank.rows[1].cells[1])
    table_bank.rows[3].cells[0].merge(table_bank.rows[3].cells[1])
    table_bank.rows[3].cells[2].merge(table_bank.rows[3].cells[3])
    
    doc.add_paragraph() # Отступ
    
    # 2. Заголовок счета
    today_str = datetime.now().strftime("%d.%m.%Y")
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"Счет на оплату № {invoice_id:05d} от {today_str} г.")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run.font.bold = True
    title_run.font.size = Pt(14)
    
    doc.add_paragraph() # Отступ
    
    # 3. Информация о сторонах (Поставщик и Покупатель)
    p_prov = doc.add_paragraph()
    p_prov.add_run("Поставщик: ").font.bold = True
    p_prov.add_run(f"{PROVIDER_REQUISITES['name']}, ИНН {PROVIDER_REQUISITES['inn']}, КПП {PROVIDER_REQUISITES['kpp']}, Адрес: {PROVIDER_REQUISITES['address']}")
    
    p_buyer = doc.add_paragraph()
    p_buyer.add_run("Покупатель: ").font.bold = True
    p_buyer.add_run(f"{tenant.full_name or tenant.name}, ИНН {tenant.inn}, КПП {tenant.kpp or '—'}, Адрес: {tenant.address or '—'}")
    
    doc.add_paragraph() # Отступ
    
    # 4. Табличная часть (Услуги)
    table_items = doc.add_table(rows=2, cols=6)
    table_items.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr = table_items.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Товары (работы, услуги)"
    hdr[2].text = "Кол-во"
    hdr[3].text = "Ед."
    hdr[4].text = "Цена"
    hdr[5].text = "Сумма"
    
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    # Сама позиция
    item_row = table_items.rows[1].cells
    item_row[0].text = "1"
    item_row[1].text = f"Лицензия на использование ПО 'СФЕРУМ-ЕРП' (Тариф '{tenant.sphere.capitalize()}', подписка 1 мес.)"
    item_row[2].text = "1"
    item_row[3].text = "мес."
    item_row[4].text = f"{amount:.2f}"
    item_row[5].text = f"{amount:.2f}"
    
    # Выравнивание
    item_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    item_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    item_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    item_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    item_row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Установка ширины колонок
    col_widths = [Inches(0.4), Inches(4.0), Inches(0.6), Inches(0.6), Inches(0.9), Inches(0.9)]
    for row in table_items.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    doc.add_paragraph() # Отступ
    
    # 5. Подвал с итогами
    p_totals = doc.add_paragraph()
    p_totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_totals.add_run(f"Итого к оплате: {amount:.2f} руб.\n").font.bold = True
    p_totals.add_run("В том числе НДС: Без НДС\n").font.italic = True
    
    # Всего наименований и сумма прописью
    sum_str = num2str_ru(amount)
    p_summary = doc.add_paragraph()
    p_summary.add_run(f"Всего наименований 1, на сумму {amount:.2f} руб.\n")
    p_summary.add_run(f"Всего к оплате: {sum_str}").font.bold = True
    
    # Разделительная линия
    doc.add_paragraph("_________________________________________________________________________________")
    
    # 6. Место для подписи / печати
    p_sign = doc.add_paragraph()
    p_sign.add_run("\nРуководитель:  __________________ (Халиков И.И.)                      Бухгалтер:  __________________ (Халикова И.И.)")
    p_sign.paragraphs[0].runs[0].font.italic = True
    
    # Сохраняем документ
    os.makedirs(output_dir, exist_ok=True)
    filename = f"invoice_{invoice_id:05d}_{uuid.uuid4().hex[:6]}.docx"
    file_path = os.path.join(output_dir, filename)
    doc.save(file_path)
    
    return file_path


def create_tenant_invoice(db: Session, tenant_id: int, amount: float = 5000.0) -> Invoice:
    """Создает запись счета в БД и запускает генерацию Docx"""
    tenant = db.query(Tenant).filter(Tenant.id == tenant_id).first()
    if not tenant:
        raise ValueError(f"Tenant with ID {tenant_id} not found")
        
    # Создаем запись счета в БД
    db_invoice = Invoice(
        tenant_id=tenant_id,
        amount=amount,
        pdf_path="generating...",  # Временная заглушка
        status="pending"
    )
    db.add(db_invoice)
    db.flush() # Получаем ID счета
    
    # Директория счетов
    backend_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    invoices_dir = os.path.join(backend_dir, "data", "invoices")
    
    # Генерируем физический файл
    try:
        file_path = generate_invoice_docx(db_invoice.id, tenant, amount, invoices_dir)
        
        # Записываем относительный путь к файлу для скачивания через веб
        relative_path = os.path.relpath(file_path, backend_dir)
        db_invoice.pdf_path = relative_path.replace("\\", "/") # Unix-style path
        db.commit()
        
        logger.info(f"[SaaS Billing] Created B2B Invoice #{db_invoice.id} for Tenant {tenant.name} ({amount} RUB)")
        return db_invoice
    except Exception as e:
        db.rollback()
        logger.error(f"[SaaS Billing] Failed to generate docx for Invoice #{db_invoice.id}: {e}")
        raise e
