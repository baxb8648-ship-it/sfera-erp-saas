from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime, timedelta
import random
import re
import os
import uuid

from ..websocket_manager import manager

from ..database import get_db
from ..models import Tender, TenderPlatform, Client, Object, User, Document, ClientStatusEnum
from ..schemas import (
    TenderCreate, TenderResponse, 
    TenderPlatformCreate, TenderPlatformResponse,
    DocumentResponse, TenderRoleResponse, TenderRoleBase
)
from .auth import get_current_user
from ..telegram import send_telegram_notification

router = APIRouter(prefix="/tenders", tags=["Tenders"])

# ----------------- TENDERS CRUD -----------------

@router.get("/", response_model=List[TenderResponse])
def get_tenders(
    status: Optional[str] = None, 
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    query = db.query(Tender)
    if status:
        query = query.filter(Tender.status == status)
    return query.order_by(Tender.submission_deadline.asc()).offset(skip).limit(limit).all()


@router.post("/", response_model=TenderResponse)
def create_tender(
    tender: TenderCreate, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    # Check for duplicates
    existing = db.query(Tender).filter(Tender.tender_number == tender.tender_number).first()
    if existing:
        raise HTTPException(
            status_code=400, 
            detail=f"Закупка с номером {tender.tender_number} уже существует в базе"
        )
        
    db_tender = Tender(**tender.model_dump())
    db.add(db_tender)
    db.commit()
    db.refresh(db_tender)
    
    from ..services.audit import log_audit_action
    log_audit_action(
        db,
        current_user,
        "создание",
        "Тендер",
        db_tender.id,
        db_tender.title,
        changes={"title": {"old": "—", "new": db_tender.title}, "tender_number": {"old": "—", "new": db_tender.tender_number}}
    )
    
    background_tasks.add_task(manager.broadcast, {

        "type": "success",
        "message": f"🆕 Добавлен новый тендер: {db_tender.title}",
        "refetchKey": "tenders"
    })
    return db_tender


@router.patch("/{tender_id}", response_model=TenderResponse)
def update_tender(
    tender_id: int, 
    tender_data: dict, # Dynamic dictionary to easily handle partial updates
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    db_tender = db.query(Tender).filter(Tender.id == tender_id).first()
    if not db_tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")
        
    old_data = {col.name: getattr(db_tender, col.name) for col in db_tender.__table__.columns}
    for key, value in tender_data.items():
        if hasattr(db_tender, key):
            setattr(db_tender, key, value)
            
    db_tender.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(db_tender)
    
    new_data = {col.name: getattr(db_tender, col.name) for col in db_tender.__table__.columns}
    from ..services.audit import get_model_changes, log_audit_action
    changes = get_model_changes(old_data, new_data)
    if changes:
        log_audit_action(db, current_user, "обновление", "Тендер", db_tender.id, db_tender.title, changes)

    
    is_won = tender_data.get("status") == "Выигран"
    msg = f"🏆 Тендер выигран! {db_tender.title}" if is_won else f"✍️ Обновлен тендер: {db_tender.title}"
    msg_type = "success" if is_won else "info"
    background_tasks.add_task(manager.broadcast, {
        "type": msg_type,
        "message": msg,
        "refetchKey": "tenders"
    })
    
    # Send Telegram notification if status has changed
    if "status" in tender_data and old_data.get("status") != tender_data["status"]:
        new_status = tender_data["status"]
        status_emojis = {
            "Анализ": "🔍",
            "Подготовка": "📝",
            "Участие": "🚀",
            "Выигран": "🏆",
            "Проигран": "❌",
            "Отказ": "⛔"
        }
        emoji = status_emojis.get(new_status, "🔄")
        tg_msg = (
            f"{emoji} <b>Статус тендера изменен:</b> {new_status}\n\n"
            f"<b>Тендер:</b> {db_tender.title}\n"
            f"<b>Заказчик:</b> {db_tender.customer_name}\n"
            f"<b>Сумма:</b> {db_tender.price:,.2f} руб."
        )
        send_telegram_notification(
            tg_msg, 
            db,
            reply_markup={
                "inline_keyboard": [
                    [{"text": "✅ Взять в работу", "callback_data": f"take_{db_tender.id}"}],
                    [
                        {"text": "👨‍💼 Назначить", "callback_data": f"assign_menu_{db_tender.id}"},
                        {"text": "📁 Документы", "url": db_tender.link if db_tender.link else "https://zakupki.gov.ru"}
                    ],
                    [{"text": "❌ Удалить (Не интересно)", "callback_data": f"delete_tender_{db_tender.id}"}]
                ]
            }
        )
        
    return db_tender


@router.delete("/{tender_id}")
def delete_tender(
    tender_id: int, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    db_tender = db.query(Tender).filter(Tender.id == tender_id).first()
    if not db_tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")
        
    tender_number = db_tender.tender_number
    # Unlink documents
    db.query(Document).filter(Document.tender_id == tender_id).update({Document.tender_id: None})
    
    # Delete associated roles
    from ..models import TenderRole
    db.query(TenderRole).filter(TenderRole.tender_id == tender_id).delete()
    
    tender_title_val = db_tender.title
    tender_id_val = db_tender.id
    db.delete(db_tender)
    db.commit()
    
    from ..services.audit import log_audit_action
    log_audit_action(
        db,
        current_user,
        "удаление",
        "Тендер",
        tender_id_val,
        tender_title_val,
        changes={"tender_number": {"old": tender_number, "new": "—"}}
    )
    
    background_tasks.add_task(manager.broadcast, {

        "type": "warning",
        "message": f"🗑️ Удален тендер №{tender_number}",
        "refetchKey": "tenders"
    })
    return {"message": "Тендер успешно удален"}


# ----------------- TENDER ROLES -----------------

@router.get("/{tender_id}/roles", response_model=List[TenderRoleResponse])
def get_tender_roles(
    tender_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from ..models import TenderRole
    return db.query(TenderRole).filter(TenderRole.tender_id == tender_id).all()

@router.post("/{tender_id}/roles", response_model=TenderRoleResponse)
def add_tender_role(
    tender_id: int,
    role_data: TenderRoleBase,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from ..models import TenderRole
    db_tender = db.query(Tender).filter(Tender.id == tender_id).first()
    if not db_tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")

    # Check if this role_name already exists for this tender
    existing_role = db.query(TenderRole).filter(
        TenderRole.tender_id == tender_id,
        TenderRole.role_name == role_data.role_name
    ).first()

    from ..telegram import send_personal_telegram_notification
    
    if existing_role:
        existing_role.user_id = role_data.user_id
        db.commit()
        db.refresh(existing_role)
        try:
            msg = (
                f"🔔 <b>Изменена роль в тендере</b>\n\n"
                f"Вы назначены на роль <b>{existing_role.role_name}</b> для тендера:\n"
                f"📌 Номер: {db_tender.tender_number}\n"
                f"📄 Название: {db_tender.title}\n"
                f"💰 Цена: {db_tender.price:,.2f} руб.\n\n"
                f"Пожалуйста, проверьте карточку закупки в CRM."
            )
            send_personal_telegram_notification(existing_role.user_id, msg, db)
        except Exception as ex:
            print(f"Failed to send update notification: {ex}")
        return existing_role
    
    new_role = TenderRole(
        tender_id=tender_id,
        user_id=role_data.user_id,
        role_name=role_data.role_name
    )
    db.add(new_role)
    db.commit()
    db.refresh(new_role)
    
    try:
        msg = (
            f"🔔 <b>Назначена роль в тендере</b>\n\n"
            f"Вы назначены на роль <b>{new_role.role_name}</b> для тендера:\n"
            f"📌 Номер: {db_tender.tender_number}\n"
            f"📄 Название: {db_tender.title}\n"
            f"💰 Цена: {db_tender.price:,.2f} руб.\n\n"
            f"Пожалуйста, проверьте карточку закупки в CRM."
        )
        send_personal_telegram_notification(new_role.user_id, msg, db)
    except Exception as ex:
        print(f"Failed to send assignment notification: {ex}")
        
    return new_role

@router.delete("/{tender_id}/roles/{role_id}")
def delete_tender_role(
    tender_id: int,
    role_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from ..models import TenderRole
    role = db.query(TenderRole).filter(TenderRole.id == role_id, TenderRole.tender_id == tender_id).first()
    if not role:
        raise HTTPException(status_code=404, detail="Роль не найдена")
    db.delete(role)
    db.commit()
    return {"message": "Роль успешно удалена"}

# ----------------- CRM CONVERSION -----------------

@router.post("/{tender_id}/participate", response_model=TenderResponse)
def participate_tender(
    tender_id: int, 
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    db_tender = db.query(Tender).filter(Tender.id == tender_id).first()
    if not db_tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")
        
    if db_tender.client_id or db_tender.object_id:
        return db_tender # Already integrated

    # 1. Create a CRM Client representing the customer
    inn_str = db_tender.inn or ""
    client_name = db_tender.customer_name or "Заказчик по тендеру"
    
    # Try finding existing client by INN first
    db_client = None
    if inn_str:
        db_client = db.query(Client).filter(Client.inn == inn_str).first()
        
    if not db_client:
        db_client = Client(
            name=client_name,
            inn=inn_str if inn_str else None,
            status=ClientStatusEnum.new,
            notes=f"Создан автоматически из тендера №{db_tender.tender_number}."
        )
        db.add(db_client)
        db.commit()
        db.refresh(db_client)

    # 2. Inquire service type from tender title/description
    service_req = "АКЗ"
    surface_type = "Металл"
    title_lower = db_tender.title.lower()
    
    if "гидроизол" in title_lower:
        service_req = "Гидроизоляция"
        surface_type = "Бетон"
    elif "пескоструй" in title_lower or "очист" in title_lower:
        service_req = "Пескоструй"
        surface_type = "Металл"
    elif "покраск" in title_lower or "окраш" in title_lower:
        service_req = "Покраска"
        surface_type = "Металл"

    # 3. Create a CRM Object
    db_object = Object(
        client_id=db_client.id,
        name=db_tender.title,
        area_sqm=1000.0, # Placeholder area
        surface_type=surface_type,
        service_required=service_req,
        status="Подготовка КП"
    )
    db.add(db_object)
    db.commit()
    db.refresh(db_object)

    # 4. Link Tender to Client and Object
    db_tender.client_id = db_client.id
    db_tender.object_id = db_object.id
    db_tender.status = "Участие"
    db_tender.updated_at = datetime.utcnow()
    
    # Link any pre-existing tender documents to the client and object
    db.query(Document).filter(Document.tender_id == tender_id).update({
        Document.client_id: db_client.id,
        Document.object_id: db_object.id
    })
    
    db.commit()
    db.refresh(db_tender)
    
    # Send TG notification about participation decision
    send_telegram_notification(
        f"🎯 <b>Принято решение об участии в тендере!</b>\n\n"
        f"📋 <b>Номер:</b> {db_tender.tender_number}\n"
        f"🏗️ <b>Проект:</b> {db_tender.title}\n"
        f"💼 <b>Заказчик:</b> {db_client.name}\n"
        f"⚙️ <b>Статус в CRM:</b> Переведен в объекты («Подготовка КП»)",
        db
    )
    
    background_tasks.add_task(manager.broadcast, {
        "type": "success",
        "message": f"🎯 Начато участие в тендере №{db_tender.tender_number}",
        "refetchKey": "tenders"
    })
    return db_tender


# ----------------- TENDER DOCUMENTS -----------------

@router.get("/{tender_id}/documents", response_model=List[DocumentResponse])
def get_tender_documents(
    tender_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    return db.query(Document).filter(Document.tender_id == tender_id).all()


@router.post("/{tender_id}/upload", response_model=DocumentResponse)
def upload_tender_document(
    tender_id: int,
    name: str = Form(...),
    doc_type: str = Form("Тендерная документация"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    db_tender = db.query(Tender).filter(Tender.id == tender_id).first()
    if not db_tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")
        
    upload_dir = os.path.join("backend", "uploads")
    os.makedirs(upload_dir, exist_ok=True)
    
    original_filename = file.filename or ""
    _, ext = os.path.splitext(original_filename)
    ext_lower = ext.lower()
    
    allowed_extensions = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".png", ".jpg", ".jpeg", ".webp"}
    if ext_lower not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail="Недопустимый формат файла. Разрешены только PDF, документы Office и изображения."
        )
        
    safe_filename = f"tender_{tender_id}_{int(datetime.utcnow().timestamp())}_{uuid.uuid4().hex}{ext_lower}"
    file_path = os.path.join(upload_dir, safe_filename)
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    content = file.file.read(MAX_FILE_SIZE + 1)
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413, 
            detail="Файл слишком большой. Максимум 50MB."
        )

    try:
        with open(file_path, "wb") as buffer:
            buffer.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
        
    new_doc = Document(
        client_id=db_tender.client_id or 0,
        object_id=db_tender.object_id,
        tender_id=tender_id,
        doc_type=doc_type,
        name=name,
        is_uploaded=1,
        file_url=file_path
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    return new_doc

@router.post("/{tender_id}/generate_doc", response_model=DocumentResponse)
def generate_tender_doc(
    tender_id: int,
    template_id: Optional[int] = Form(None),
    doc_type: str = Form("Заявка на участие"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from ..services.document_generator import generate_tender_document
    from docx import Document as DocxDocument
    from ..models import DocumentTemplate
    
    db_tender = db.query(Tender).filter(Tender.id == tender_id).first()
    if not db_tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")
        
    templates_dir = os.path.join("backend", "templates")
    os.makedirs(templates_dir, exist_ok=True)
    
    template_path = None
    if template_id:
        db_template = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id, DocumentTemplate.is_active == 1).first()
        if db_template:
            template_path = db_template.file_path
            doc_type = db_template.doc_type
            
    if not template_path:
        # Fallback to default template
        template_name = "application_template.docx"
        template_path = os.path.join(templates_dir, template_name)
        
        # Create a dummy template if it doesn't exist for demo purposes
        if not os.path.exists(template_path):
            dummy = DocxDocument()
            dummy.add_heading('ЗАЯВКА НА УЧАСТИЕ В ТЕНДЕРЕ', 0)
            dummy.add_paragraph('Тендер: {{TENDER_TITLE}}')
            dummy.add_paragraph('Номер: {{TENDER_NUMBER}}')
            dummy.add_paragraph('Заказчик: {{CUSTOMER_NAME}}')
            dummy.add_paragraph('Сумма (НМЦК): {{PRICE}} руб.')
            dummy.add_paragraph('Дата подачи: {{CURRENT_DATE}}')
            dummy.add_paragraph('Исполнитель: ООО "СФЕРУМ"')
            dummy.save(template_path)
        
    output_dir = os.path.join("backend", "uploads")
    
    data = {
        "TENDER_TITLE": db_tender.title,
        "TENDER_NUMBER": db_tender.tender_number,
        "CUSTOMER_NAME": db_tender.customer_name or "Не указан",
        "PRICE": f"{db_tender.price:,.2f}" if db_tender.price else "0.00",
        "CURRENT_DATE": datetime.utcnow().strftime("%d.%m.%Y")
    }
    
    try:
        output_path = generate_tender_document(
            template_path=template_path,
            output_dir=output_dir,
            data=data,
            filename_prefix="zayavka"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Doc generation failed: {str(e)}")
        
    # Relative path for URL
    rel_path = os.path.relpath(output_path, "backend")
    # Normalize slashes for URL
    rel_path = rel_path.replace("\\", "/")
    
    new_doc = Document(
        client_id=db_tender.client_id or 0,
        object_id=db_tender.object_id,
        tender_id=tender_id,
        doc_type=doc_type,
        name=f"Авто-документ: {doc_type}",
        is_uploaded=1,
        file_url=rel_path
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    
    return new_doc


# ----------------- DOCUMENT TEMPLATES API -----------------
from ..schemas import DocumentTemplateResponse

@router.get("/templates", response_model=List[DocumentTemplateResponse])
def get_document_templates(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from ..models import DocumentTemplate
    return db.query(DocumentTemplate).filter(DocumentTemplate.is_active == 1).all()

@router.post("/templates/upload", response_model=DocumentTemplateResponse)
def upload_document_template(
    name: str = Form(...),
    doc_type: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from ..models import DocumentTemplate
    templates_dir = os.path.join("backend", "templates")
    os.makedirs(templates_dir, exist_ok=True)
    
    ext = os.path.splitext(file.filename)[1]
    if ext.lower() != ".docx":
        raise HTTPException(status_code=400, detail="Разрешены только файлы формата .docx")
        
    safe_filename = f"template_{int(datetime.utcnow().timestamp())}_{uuid.uuid4().hex[:6]}.docx"
    file_path = os.path.join(templates_dir, safe_filename)
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    content = file.file.read(MAX_FILE_SIZE + 1)
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413, 
            detail="Файл слишком большой. Максимум 50MB."
        )

    try:
        with open(file_path, "wb") as buffer:
            buffer.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Не удалось сохранить шаблон: {str(e)}")
        
    db_template = DocumentTemplate(
        name=name,
        doc_type=doc_type,
        file_path=file_path,
        is_active=1
    )
    db.add(db_template)
    db.commit()
    db.refresh(db_template)
    
    return db_template

@router.delete("/templates/{template_id}")
def delete_document_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    from ..models import DocumentTemplate
    db_template = db.query(DocumentTemplate).filter(DocumentTemplate.id == template_id).first()
    if not db_template:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
        
    db_template.is_active = 0
    db.commit()
    return {"message": "Шаблон удален"}


# ----------------- MONITORS / PLATFORMS API -----------------

@router.get("/platforms", response_model=List[TenderPlatformResponse])
def get_platforms(
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    return db.query(TenderPlatform).all()


@router.post("/platforms", response_model=TenderPlatformResponse)
def create_or_update_platform(
    platform: TenderPlatformCreate, 
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    # Check if exists by name
    db_platform = db.query(TenderPlatform).filter(TenderPlatform.name == platform.name).first()
    if db_platform:
        # Update existing
        for key, value in platform.model_dump().items():
            setattr(db_platform, key, value)
    else:
        # Create new
        db_platform = TenderPlatform(**platform.model_dump())
        db.add(db_platform)
        
    db.commit()
    db.refresh(db_platform)
    return db_platform


@router.delete("/platforms/{platform_id}")
def delete_platform(
    platform_id: int, 
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    db_platform = db.query(TenderPlatform).filter(TenderPlatform.id == platform_id).first()
    if not db_platform:
        raise HTTPException(status_code=404, detail="Площадка не найдена")
    db.delete(db_platform)
    db.commit()
    return {"message": "Площадка успешно удалена"}


@router.post("/sync")
def sync_tenders(
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db), 
    current_user: User = Depends(get_current_user)
):
    from ..models import CompanySetting
    sync_mode_setting = db.query(CompanySetting).filter(CompanySetting.key == "tender_sync_mode").first()
    sync_mode = sync_mode_setting.value if sync_mode_setting else "demo"

    if sync_mode == "live":
        active_platforms = db.query(TenderPlatform).filter(TenderPlatform.is_active == 1).all()
        if not active_platforms:
            return {"message": "Синхронизация завершена. Нет активных настроек площадок.", "imported_count": 0}

        import urllib.request
        import urllib.parse
        import xml.etree.ElementTree as ET
        import ssl
        
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "ru-RU,ru;q=0.8,en-US;q=0.5,en;q=0.3"
        }

        imported_tenders = []
        imported_numbers = set()
        
        for platform in active_platforms:
            is_zakupki = "zakupki" in platform.name.lower() or "zakupki" in platform.api_url.lower()
            is_b2b = "b2b" in platform.name.lower() or "b2b" in platform.api_url.lower()
            is_fabrikant = "fabrikant" in platform.name.lower() or "fabrikant" in platform.api_url.lower() or "фабрикант" in platform.name.lower()
            is_roseltorg = "roseltorg" in platform.name.lower() or "roseltorg" in platform.api_url.lower() or "росэлторг" in platform.name.lower()
            is_tektorg = "tektorg" in platform.name.lower() or "tektorg" in platform.api_url.lower() or "тэк-торг" in platform.name.lower() or "тэкторг" in platform.name.lower()
            is_etpgpb = "etpgpb" in platform.name.lower() or "etpgpb" in platform.api_url.lower() or "гпб" in platform.name.lower() or "газпромбанк" in platform.name.lower()
            is_tenderpro = "tender.pro" in platform.name.lower() or "tenderpro" in platform.name.lower() or "тендерпро" in platform.name.lower() or "tender.pro" in platform.api_url.lower()

            keywords_list = [k.strip() for k in re.split(r'[,; ]+', platform.keywords) if k.strip()]
            regions_list = [r.strip() for r in re.split(r'[,;]+', platform.regions) if r.strip()]
            
            exclude_list = []
            if platform.exclude_keywords:
                exclude_list = [ek.strip().lower() for ek in re.split(r'[,; ]+', platform.exclude_keywords) if ek.strip()]

            parser_instance = None
            if is_b2b:
                from ..parsers.b2b_center import B2BCenterParser
                parser_instance = B2BCenterParser()
            elif is_fabrikant:
                from ..parsers.fabrikant import FabrikantParser
                parser_instance = FabrikantParser()
            elif is_roseltorg:
                from ..parsers.roseltorg import RoseltorgParser
                parser_instance = RoseltorgParser()
            elif is_tektorg:
                from ..parsers.tektorg import TEKTorgParser
                parser_instance = TEKTorgParser()
            elif is_etpgpb:
                from ..parsers.etpgpb import ETPGPBParser
                parser_instance = ETPGPBParser()
            elif is_tenderpro:
                from ..parsers.tenderpro import TenderProParser
                parser_instance = TenderProParser()

            if parser_instance:
                try:
                    from ..utils.customer_analyzer import analyze_customer
                    parsed_data = parser_instance.parse(platform, imported_numbers)
                    for item in parsed_data:
                        analytics = analyze_customer(item.get("inn"), platform.name)
                        item["expected_dumping"] = analytics["expected_dumping"]
                        item["expected_participants"] = analytics["expected_participants"]
                        new_tender = Tender(**item)
                        db.add(new_tender)
                        imported_tenders.append(new_tender)
                        
                        if len(imported_tenders) >= 10:
                            break
                    platform.last_sync = datetime.utcnow()
                    db.commit()
                except Exception as e:
                    print(f"Error executing parser for {platform.name}: {e}")
                continue
                
            if not is_zakupki:
                continue
            
            # Exclude keywords list
            exclude_list = []
            if platform.exclude_keywords:
                exclude_list = [ek.strip().lower() for ek in re.split(r'[,; ]+', platform.exclude_keywords) if ek.strip()]

            # Fetch tenders for each keyword
            for keyword in keywords_list:
                encoded_keyword = urllib.parse.quote(keyword)
                url = f"https://zakupki.gov.ru/epz/order/extendedsearch/rss.html?searchString={encoded_keyword}"
                
                try:
                    req = urllib.request.Request(url, headers=headers)
                    with urllib.request.urlopen(req, context=ctx, timeout=15) as response:
                        content = response.read()
                        
                    root = ET.fromstring(content)
                    items = root.findall('.//item')
                    
                    for item in items[:15]:
                        link_url = item.find('link').text if item.find('link') is not None else ""
                        pub_date_str = item.find('pubDate').text if item.find('pubDate') is not None else ""
                        description_html = item.find('description').text if item.find('description') is not None else ""
                        
                        # Extract regNumber from link or title
                        reg_number_match = re.search(r'regNumber=(\d+)', link_url)
                        if not reg_number_match:
                            continue
                        tender_num = reg_number_match.group(1)
                        
                        # Check if tender already exists in DB
                        if tender_num in imported_numbers:
                            continue
                        exists = db.query(Tender).filter(Tender.tender_number == tender_num).first()
                        if exists:
                            continue
                            
                        title_match = re.search(r'Наименование объекта закупки:.*?(?:</strong>|&lt;/strong&gt;|)\s*(.*?)(?:<br/?>|&lt;br/?&gt;|$)', description_html, re.IGNORECASE)
                        if not title_match:
                            title_match = re.search(r'Наименование объекта закупки:\s*(.*?)(?:<br/?>|$)', description_html, re.IGNORECASE)
                        
                        tender_title = title_match.group(1).strip() if title_match else ""
                        tender_title = re.sub(r'<[^>]*>', '', tender_title).strip()
                        
                        if not tender_title or tender_title.lower() in ("null", "none", ""):
                            raw_item_title = item.find('title').text if item.find('title') is not None else ""
                            tender_title = raw_item_title if raw_item_title else "Без названия"
                        
                        # Customer
                        cust_match = re.search(r'Наименование Заказчика:.*?(?:</strong>|&lt;/strong&gt;|)\s*(.*?)(?:<br/?>|&lt;br/?&gt;|$)', description_html, re.IGNORECASE)
                        if not cust_match:
                            cust_match = re.search(r'Наименование Заказчика:\s*(.*?)(?:<br/?>|$)', description_html, re.IGNORECASE)
                        cust_name = cust_match.group(1).strip() if cust_match else "Не указан"
                        cust_name = re.sub(r'<[^>]*>', '', cust_name).strip()
                        
                        # Price
                        price_match = re.search(r'Начальная цена контракта:.*?(?:</strong>|&lt;/strong&gt;)\s*([\d\.\s\xa0]+)', description_html, re.IGNORECASE)
                        if not price_match:
                            price_match = re.search(r'Начальная цена контракта:\s*([\d\.\s\xa0]+)', description_html, re.IGNORECASE)
                        
                        price = 0.0
                        if price_match:
                            price_str = price_match.group(1).replace(" ", "").replace("\xa0", "").strip()
                            try:
                                price = float(price_str)
                            except ValueError:
                                price = 0.0
                        
                        # Exclude keywords filter
                        if exclude_list:
                            if any(ek in tender_title.lower() or ek in description_html.lower() for ek in exclude_list):
                                continue
                                
                        # Price limits filter
                        if platform.min_price and price < platform.min_price:
                            continue
                        if platform.max_price and price > platform.max_price:
                            continue
                        
                        # Publication date
                        pub_date = datetime.utcnow()
                        if pub_date_str:
                            for date_fmt in ('%a, %d %b %Y %H:%M:%S %Z', '%a, %d %b %Y %H:%M:%S %z'):
                                try:
                                    pub_date = datetime.strptime(pub_date_str, date_fmt)
                                    break
                                except ValueError:
                                    continue
                        
                        # Load HTML details for INN, deadline, and region
                        inn = None
                        deadline = None
                        region_match_found = False if regions_list else True
                        
                        detail_html = ""
                        extracted_volumes = ""
                        
                        if link_url:
                            try:
                                # Fetch common info for INN, deadline
                                req_detail = urllib.request.Request(link_url, headers=headers)
                                with urllib.request.urlopen(req_detail, context=ctx, timeout=10) as resp_detail:
                                    detail_html = resp_detail.read().decode('utf-8', errors='ignore')
                                    
                                # Fetch documents page for PDF/Word parsing
                                if "regNumber=" in link_url:
                                    docs_url = link_url.replace("common-info.html", "documents.html")
                                    try:
                                        req_docs = urllib.request.Request(docs_url, headers=headers)
                                        with urllib.request.urlopen(req_docs, context=ctx, timeout=10) as resp_docs:
                                            docs_html = resp_docs.read().decode('utf-8', errors='ignore')
                                            
                                        # Find PDF or DOCX links
                                        doc_links = re.findall(r'href="([^"]+\.(?:pdf|docx|doc)[^"]*)"', docs_html, re.IGNORECASE)
                                        if doc_links:
                                            from ..utils.document_parser import parse_document_from_url
                                            for dlink in doc_links[:2]: # check first 2 files
                                                full_dlink = dlink if dlink.startswith("http") else "https://zakupki.gov.ru" + dlink
                                                ext = "pdf" if ".pdf" in dlink.lower() else "docx"
                                                volumes = parse_document_from_url(full_dlink, ext)
                                                if volumes:
                                                    extracted_volumes += volumes + "\n"
                                                    break # stop after finding volumes in one file
                                    except Exception as ed:
                                        print(f"Error fetching docs: {ed}")
                                        
                            except Exception as ex:
                                pass
                                
                        if detail_html:
                            # 1. Parse INN
                            inn_match = re.search(r'inn=(\d{10,12})', detail_html)
                            if inn_match:
                                inn = inn_match.group(1)
                            else:
                                text_inn_match = re.search(r'(?:ИНН|inn)\s*[\:\-]?\s*(\d{10,12})', detail_html, re.IGNORECASE)
                                if text_inn_match:
                                    inn = text_inn_match.group(1)
                                else:
                                    potential_inns = re.findall(r'\b\d{10}\b|\b\d{12}\b', detail_html)
                                    for p in potential_inns:
                                        if p not in ('1457694118', '200000000000'):
                                            inn = p
                                            break
                            
                            # 2. Parse Deadline
                            p1 = re.search(r'Окончание подачи заявок.*?class="data-block__value"\s*>\s*([\d\.\s\:]+)\s*</div>', detail_html, re.DOTALL | re.IGNORECASE)
                            deadline_str = p1.group(1).strip() if p1 else None
                            if not deadline_str:
                                p2 = re.search(r'Дата и время окончания срока подачи заявок.*?class="common-text__value"\s*>\s*([\d\.\s\:]+)\s*</div>', detail_html, re.DOTALL | re.IGNORECASE)
                                deadline_str = p2.group(1).strip() if p2 else None
                                
                            if deadline_str:
                                deadline_clean = re.sub(r'\s+', ' ', deadline_str).strip()
                                for fmt in ('%d.%m.%Y %H:%M', '%d.%m.%Y'):
                                    try:
                                        deadline = datetime.strptime(deadline_clean, fmt)
                                        break
                                    except ValueError:
                                        continue
                                        
                            # 3. Check regions filter
                            if regions_list:
                                for r in regions_list:
                                    if r.lower() in detail_html.lower():
                                        region_match_found = True
                                        break
                                        
                        # Убрана фильтрация минус-слов по сырому HTML, так как теперь работает умный AI-фильтр

                        # Fallback for missing fields
                        if not deadline:
                            deadline = pub_date + timedelta(days=7)
                            
                        # Regions filter verification
                        if regions_list and not region_match_found:
                            # Search in title or description_html as backup
                            for r in regions_list:
                                if r.lower() in tender_title.lower() or r.lower() in description_html.lower():
                                    region_match_found = True
                                    break
                            if not region_match_found:
                                continue
                                
                        # Add extracted volumes to description if found
                        final_description = description_html[:2000]
                        
                        from ..utils.ai_engine import ai_filter_tender, ai_summarize_tender
                        
                        # Пункт 5: Умный AI Фильтр
                        if not ai_filter_tender(tender_title, final_description):
                            # AI решил, что это мусор (например, поставка материалов)
                            print(f"Skipped {tender_num} due to AI Filter")
                            continue
                            
                        # Пункт 3: AI Саммаризация документов
                        if extracted_volumes:
                            summary = ai_summarize_tender(extracted_volumes)
                            if summary:
                                final_description = f"<b>[AI-СВОДКА ТЗ]:</b><br/>{summary}<br/><br/>" + final_description
                            else:
                                final_description = f"<b>[ИЗВЛЕЧЕНО ИЗ ДОКУМЕНТАЦИИ ТЗ]:</b><br/>{extracted_volumes[:1000]}<br/><br/>" + final_description

                        from ..utils.customer_analyzer import analyze_customer
                        analytics = analyze_customer(inn, "Закупки.gov.ru")

                        # Instantiate Tender
                        new_tender = Tender(
                            tender_number=tender_num,
                            title=tender_title,
                            description=final_description,
                            customer_name=cust_name,
                            inn=inn,
                            price=price,
                            currency="RUB",
                            platform="Закупки.gov.ru",
                            link=link_url,
                            status="Анализ",
                            publication_date=pub_date,
                            submission_deadline=deadline,
                            expected_dumping=analytics["expected_dumping"],
                            expected_participants=analytics["expected_participants"]
                        )
                        
                        db.add(new_tender)
                        imported_tenders.append(new_tender)
                        imported_numbers.add(tender_num)
                        
                        # Max 10 per sync to avoid spam / overloading
                        if len(imported_tenders) >= 10:
                            break
                    
                except Exception as ex:
                    print(f"Error syncing keyword '{keyword}': {ex}")
                    continue
                    
                if len(imported_tenders) >= 10:
                    break
                    
            platform.last_sync = datetime.utcnow()
            
        db.commit()
        
        # Send Telegram alerts for imported tenders
        if imported_tenders:
            for tender in imported_tenders:
                db.refresh(tender)
                deadline_str = tender.submission_deadline.strftime('%d.%m.%Y %H:%M') if tender.submission_deadline else "Не указано"
                
                # Check if description has extracted volumes
                extracted_info = ""
                if "<b>[AI-СВОДКА ТЗ]:</b>" in tender.description:
                    ext_parts = tender.description.split("<br/><br/>")
                    if ext_parts:
                        clean_ext = ext_parts[0].replace("<b>[AI-СВОДКА ТЗ]:</b><br/>", "").replace("<br/>", "\n")
                        extracted_info = f"\n🤖 <b>AI-Выжимка из документов:</b>\n<i>{clean_ext}</i>\n"
                elif "<b>[ИЗВЛЕЧЕНО ИЗ ДОКУМЕНТАЦИИ ТЗ]:</b>" in tender.description:
                    ext_parts = tender.description.split("<br/><br/>")
                    if ext_parts:
                        clean_ext = ext_parts[0].replace("<b>[ИЗВЛЕЧЕНО ИЗ ДОКУМЕНТАЦИИ ТЗ]:</b><br/>", "").replace("<br/>", "\n")
                        extracted_info = f"\n📄 <b>Извлечено из ТЗ:</b>\n<i>{clean_ext}</i>\n"

                # Analytics string
                analytics_info = ""
                if tender.expected_dumping and tender.expected_participants:
                    analytics_info = (
                        f"\n📈 <b>Аналитика Заказчика (Прогноз):</b>\n"
                        f"📉 Демпинг: <b>{tender.expected_dumping}</b>\n"
                        f"👥 Конкуренты: <b>{tender.expected_participants}</b>\n"
                    )

                send_telegram_notification(
                    f"📢 <b>Найден новый подходящий тендер! (LIVE)</b>\n\n"
                    f"🔍 <b>Название:</b> {tender.title}\n"
                    f"💰 <b>НМЦК:</b> {tender.price:,.2f} руб.\n"
                    f"🏢 <b>Заказчик:</b> {tender.customer_name} (ИНН: {tender.inn})\n"
                    f"📍 <b>Площадка:</b> {tender.platform}\n"
                    f"📅 <b>Подача заявок до:</b> {deadline_str}\n"
                    f"{extracted_info}"
                    f"{analytics_info}"
                    f"🔗 <a href='{tender.link}'>Ссылка на тендер</a>",
                    db,
                    reply_markup={
                        "inline_keyboard": [
                            [{"text": "✅ Взять в работу", "callback_data": f"take_{tender.id}"}],
                            [
                                {"text": "👨‍💼 Назначить", "callback_data": f"assign_menu_{tender.id}"},
                                {"text": "🤖 ИИ-Анализ", "callback_data": f"ai_analyze_{tender.id}"}
                            ],
                            [{"text": "📁 Документы", "url": tender.link}] if tender.link else [],
                            [{"text": "❌ Удалить (Не интересно)", "callback_data": f"delete_tender_{tender.id}"}]
                        ]
                    }
                )
                
        return {
            "message": f"Синхронизация успешно завершена (Боевой режим). Импортировано {len(imported_tenders)} новых тендеров.",
            "imported_count": len(imported_tenders)
        }


    active_platforms = db.query(TenderPlatform).filter(TenderPlatform.is_active == 1).all()
    if not active_platforms:
        return {"message": "Синхронизация завершена. Нет активных API настроек площадок.", "imported_count": 0}

    # Pre-defined realistic tender titles to match keywords
    tender_templates = [
        {"keyword_match": "антикор", "title": "Выполнение работ по антикоррозийной защите резервуаров вертикальных стальных РВС-20000", "desc": "Очистка пескоструйным методом и антикоррозийная покраска металлических резервуаров нефтебазы."},
        {"keyword_match": "гидроизол", "title": "Гидроизоляция железобетонных конструкций резервуаров чистой воды (РЧВ)", "desc": "Устройство проникающей и обмазочной гидроизоляции бетонных стен и перекрытий очистных сооружений."},
        {"keyword_match": "покрас", "title": "Окрашивание металлических пролетных строений и опор путепровода автомобильной дороги", "desc": "Подготовка поверхностей, обеспыливание, грунтование и финишная покраска антикоррозийными красками."},
        {"keyword_match": "огнезащит", "title": "Нанесение огнезащитного и антикоррозийного покрытия на несущие металлоконструкции производственного корпуса", "desc": "Подготовка металлических балок и нанесение огнезащитной вспучивающейся краски."},
        {"keyword_match": "пескоструй", "title": "Абразивоструйная очистка и восстановление защитного слоя железобетонных плит перекрытия эстакады", "desc": "Очистка бетонных поверхностей от разрушившихся слоев пескоструем с последующим торкретированием."},
        {"keyword_match": "антикор", "title": "Работы по АКЗ трубопроводов технологической эстакады на объекте ПАО 'Транснефть'", "desc": "Комплекс окрасочных антикоррозийных работ на наружных технологических трубопроводах."},
        {"keyword_match": "покрас", "title": "Косметический ремонт и покраска фасада здания административно-бытового комплекса", "desc": "Шпатлевание, грунтование и окраска фасадными акриловыми красками высокой стойкости."}
    ]

    customers = [
        {"name": "ПАО 'Газпром нефть'", "inn": "7706596240"},
        {"name": "АО 'Оренбургнефть'", "inn": "5603002235"},
        {"name": "Министерство транспорта и дорожного хозяйства Республики Башкортостан", "inn": "0274154944"},
        {"name": "ООО 'Уфимский нефтеперерабатывающий завод'", "inn": "0282006540"},
        {"name": "Администрация городского округа Самара", "inn": "6315800124"},
        {"name": "АО 'Транснефть - Приволга'", "inn": "6316002620"}
    ]

    imported_tenders = []
    
    for platform in active_platforms:
        # Split keywords by commas, spaces
        keywords_list = [k.strip().lower() for k in re.split(r'[,; ]+', platform.keywords) if k.strip()]
        regions_list = [r.strip() for r in re.split(r'[,;]+', platform.regions) if r.strip()]
        
        # If keywords are empty, default to broad search
        if not keywords_list:
            keywords_list = ["антикор"]

        # Loop through templates to find matches
        for template in tender_templates:
            # Check if template match matches any search keyword
            matches_keyword = any(kw in template["title"].lower() or kw in template["desc"].lower() or kw in template["keyword_match"] for kw in keywords_list)
            if not matches_keyword:
                continue

            # Check if template matches any exclude keyword (minus-words)
            if platform.exclude_keywords:
                exclude_list = [ek.strip().lower() for ek in re.split(r'[,; ]+', platform.exclude_keywords) if ek.strip()]
                if any(ek in template["title"].lower() or ek in template["desc"].lower() for ek in exclude_list):
                    continue

            # Generate random tender details
            rand_suffix = str(random.randint(100000, 999999))
            if "eis" in platform.api_url.lower() or "zakupki" in platform.name.lower():
                tender_num = "01732000014240" + rand_suffix
                link_url = f"https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber={tender_num}"
                plat_name = "Закупки.gov.ru"
            else:
                tender_num = "B2B-" + rand_suffix
                link_url = f"https://www.b2b-center.ru/market/tender-{rand_suffix}/"
                plat_name = platform.name

            # Check if this tender is already in our DB
            exists = db.query(Tender).filter(Tender.tender_number == tender_num).first()
            if exists:
                continue

            # Check price boundaries
            price = float(random.randint(2000, 15000) * 1000) # 2.0M to 15.0M
            if platform.min_price and price < platform.min_price:
                price = platform.min_price + 100000
            if platform.max_price and price > platform.max_price:
                price = platform.max_price - 100000

            # Pick customer
            cust = random.choice(customers)
            
            # Select region
            region_selected = "Оренбургская область"
            if regions_list:
                region_selected = random.choice(regions_list)

            # Deadline (5 to 15 days in future)
            deadline = datetime.utcnow() + timedelta(days=random.randint(5, 15), hours=random.randint(0, 23))

            # Instantiate Tender
            new_tender = Tender(
                tender_number=tender_num,
                title=f"[ДЕМО] {template['title']} ({region_selected})",
                description=template["desc"],
                customer_name=cust["name"],
                inn=cust["inn"],
                price=price,
                currency="RUB",
                platform=plat_name,
                link=link_url,
                status="Анализ",
                publication_date=datetime.utcnow() - timedelta(days=random.randint(1, 3)),
                submission_deadline=deadline
            )
            
            db.add(new_tender)
            imported_tenders.append(new_tender)
            
            # Break if we imported enough in this sync cycle (max 3 per sync to keep it realistic)
            if len(imported_tenders) >= 3:
                break
                
        platform.last_sync = datetime.utcnow()

    db.commit()

    # Send Telegram alerts for imported tenders
    if imported_tenders:
        for tender in imported_tenders:
            db.refresh(tender)
            send_telegram_notification(
                f"📢 <b>Найден новый подходящий тендер!</b>\n\n"
                f"🔍 <b>Название:</b> {tender.title}\n"
                f"💰 <b>НМЦК:</b> {tender.price:,.2f} руб.\n"
                f"🏢 <b>Заказчик:</b> {tender.customer_name} (ИНН: {tender.inn})\n"
                f"📍 <b>Площадка:</b> {tender.platform}\n"
                f"📅 <b>Подача заявок до:</b> {tender.submission_deadline.strftime('%d.%m.%Y %H:%M')}\n"
                f"🔗 <a href='{tender.link}'>Ссылка на тендер</a>",
                db,
                reply_markup={
                    "inline_keyboard": [
                        [{"text": "✅ Взять в работу", "callback_data": f"take_{tender.id}"}],
                        [
                            {"text": "👨‍💼 Назначить", "callback_data": f"assign_menu_{tender.id}"},
                            {"text": "🤖 ИИ-Анализ", "callback_data": f"ai_analyze_{tender.id}"}
                        ],
                        [{"text": "📁 Документы", "url": tender.link}] if tender.link else [],
                        [{"text": "❌ Удалить (Не интересно)", "callback_data": f"delete_tender_{tender.id}"}]
                    ]
                }
            )

    return {
        "message": f"Синхронизация успешно завершена (Демо-режим). Импортировано {len(imported_tenders)} новых тендеров.",
        "imported_count": len(imported_tenders)
    }

@router.post("/{tender_id}/ai-analyze", response_model=TenderResponse)
def analyze_tender_ai(
    tender_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    import re
    from ..utils.ai_engine import ask_ollama

    db_tender = db.query(Tender).filter(Tender.id == tender_id).first()
    if not db_tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")

    clean_desc = re.sub(r'<[^>]*>', '', db_tender.description or "")[:3000]
    prompt = (
        "Ты — ИИ-Копилот, технический эксперт компании ООО СФЕРУМ по антикоррозийной защите (АКЗ), "
        "огнезащите металлоконструкций, подготовке поверхностей (Sa 2.5, Sa 3, ГОСТ 9.402) и покраске ЛКМ.\n"
        "Проведи краткий экспресс-анализ этой закупки.\n\n"
        f"Название: {db_tender.title}\n"
        f"Заказчик: {db_tender.customer_name}\n"
        f"НМЦК: {db_tender.price:,.2f} руб.\n"
        f"ТЗ/Описание: {clean_desc}\n\n"
        "Напиши структурированный технический экспресс-анализ на русском языке:\n"
        "1. 🎯 Суть проекта (кратко)\n"
        "2. 🛠️ Ключевые объемы и требования к подготовке/ЛКМ\n"
        "3. ⚠️ Сложности и риски (высота, стесненность, погодные условия, сжатые сроки)\n"
        "4. 📊 Прогноз цены (учитывая демпинг заказчика: " + (db_tender.expected_dumping or "неизвестно") + ")\n\n"
        "Отвечай емко, по существу, профессионально и без лишних вводных слов."
    )

    ai_response = ask_ollama(prompt)
    if not ai_response:
        raise HTTPException(
            status_code=503,
            detail="Локальный сервер ИИ (Ollama) временно недоступен или модель qwen2:7b не загружена"
        )

    db_tender.ai_analysis = ai_response
    db.commit()
    db.refresh(db_tender)
    return db_tender

